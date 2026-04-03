from __future__ import annotations

import re
import warnings
from copy import copy
from pathlib import Path
from tempfile import NamedTemporaryFile
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile

import pandas as pd
import plotly.express as px
from openpyxl import load_workbook

from .analysis import DECAY_METHODS
from .configuration import ConfigManager
from .models import AnalysisArtifacts, AnalysisPlan, AnalysisResult, WorkbookContext
from .utils import ensure_directory, jsonify, normalize_text, safe_filename, timestamp_slug, write_json

try:
    from docx import Document
except ImportError:  # pragma: no cover - optional dependency
    Document = None

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfgen import canvas
except ImportError:  # pragma: no cover - optional dependency
    A4 = None
    canvas = None
    pdfmetrics = None
    UnicodeCIDFont = None


class ReportGeneratorAgent:
    """Persist analysis artifacts in multiple formats."""

    def __init__(self, output_root: str | Path, config_manager: ConfigManager | None = None) -> None:
        self.output_root = ensure_directory(Path(output_root).resolve())
        self.config_manager = config_manager
        self.business_rules = config_manager.business_rules() if config_manager else {}

    def generate(
        self,
        workbook: WorkbookContext,
        analysis_plan: AnalysisPlan,
        analysis_result: AnalysisResult,
        template_file: str | Path | None = None,
    ) -> AnalysisArtifacts:
        if analysis_result.summary_frame is None:
            raise ValueError("缺少 summary_frame，无法生成输出文件。")

        job_name = workbook.file_path.stem if workbook.file_path else "analysis"
        job_dir = ensure_directory(self.output_root / f"{timestamp_slug()}_{safe_filename(job_name)}")
        artifacts = AnalysisArtifacts(output_dir=job_dir)

        excel_path = job_dir / "analysis_output.xlsx"
        template_path = Path(template_file).resolve() if template_file else workbook.template_file
        if template_path:
            try:
                self._write_excel_with_template(excel_path, template_path, analysis_plan, analysis_result)
            except Exception as exc:
                analysis_result.warnings.append(f"模板化导出失败，已回退为标准 Excel 导出：{exc}")
                self._write_excel_standard(excel_path, analysis_plan, analysis_result)
        else:
            self._write_excel_standard(excel_path, analysis_plan, analysis_result)
        artifacts.excel_path = excel_path

        if "chart" in analysis_plan.output_formats:
            artifacts.chart_paths = self._write_charts(job_dir, analysis_result)

        markdown_text = self._build_markdown_report(workbook, analysis_plan, analysis_result, artifacts)
        markdown_path = job_dir / "analysis_report.md"
        markdown_path.write_text(markdown_text, encoding="utf-8")
        artifacts.markdown_path = markdown_path

        audit_report_path = job_dir / "audit_report.md"
        audit_report_path.write_text(
            self._build_audit_report(workbook, analysis_plan, analysis_result, artifacts),
            encoding="utf-8",
        )
        artifacts.audit_report_path = audit_report_path

        if "docx" in analysis_plan.output_formats:
            docx_path = job_dir / "analysis_report.docx"
            if self._write_docx(docx_path, markdown_text):
                artifacts.docx_path = docx_path
            else:
                analysis_result.warnings.append("当前环境缺少 python-docx，已跳过 Docx 报告导出。")

        if "pdf" in analysis_plan.output_formats:
            pdf_path = job_dir / "analysis_report.pdf"
            if self._write_pdf(pdf_path, markdown_text):
                artifacts.pdf_path = pdf_path
            else:
                analysis_result.warnings.append("当前环境缺少 reportlab，已跳过 PDF 报告导出。")

        audit_log_path = job_dir / "audit_log.json"
        write_json(
            audit_log_path,
            {
                "workbook": jsonify(workbook),
                "analysis_plan": jsonify(analysis_plan),
                "analysis": {
                    "source_sheet": analysis_result.source_sheet,
                    "source_files": analysis_result.source_files,
                    "base_field_mapping": analysis_result.base_field_mapping,
                    "metric_matches": jsonify(analysis_result.metric_matches),
                    "warnings": analysis_result.warnings,
                    "validation_messages": analysis_result.validation_messages,
                    "narrative_summary": analysis_result.narrative_summary,
                    "llm_notes": analysis_result.llm_notes,
                    "statistics": analysis_result.statistics,
                    "comparison_notes": analysis_result.comparison_notes,
                    "audit_records": analysis_result.audit_records,
                    "recompute_reviews": analysis_result.recompute_reviews,
                    "trace_events": jsonify(analysis_result.trace_events),
                },
            },
        )
        artifacts.audit_log_path = audit_log_path
        return artifacts

    def _write_excel_standard(
        self,
        excel_path: Path,
        analysis_plan: AnalysisPlan,
        analysis_result: AnalysisResult,
    ) -> None:
        summary_frame = analysis_result.summary_frame
        assert summary_frame is not None

        metric_match_frame = pd.DataFrame([item.model_dump(mode="json") for item in analysis_result.metric_matches])
        statistics_frame = pd.DataFrame.from_dict(analysis_result.statistics, orient="index")
        statistics_frame.index.name = "metric"
        validation_frame = pd.DataFrame(
            {"message": analysis_result.validation_messages + analysis_result.warnings + analysis_result.comparison_notes}
        )
        audit_frame = pd.DataFrame(analysis_result.audit_records)

        with pd.ExcelWriter(excel_path, engine="openpyxl") as writer:
            summary_frame.to_excel(writer, sheet_name=analysis_plan.output_sheet_name[:31], index=False)
            metric_match_frame.to_excel(writer, sheet_name="匹配结果", index=False)
            statistics_frame.to_excel(writer, sheet_name="统计汇总")
            validation_frame.to_excel(writer, sheet_name="审计告警", index=False)
            audit_frame.to_excel(writer, sheet_name="审计明细", index=False)

    def _write_excel_with_template(
        self,
        excel_path: Path,
        template_path: Path,
        analysis_plan: AnalysisPlan,
        analysis_result: AnalysisResult,
    ) -> None:
        summary_frame = analysis_result.summary_frame
        assert summary_frame is not None

        with warnings.catch_warnings():
            warnings.filterwarnings("ignore", message="wmf image format is not supported.*", category=UserWarning)
            warnings.filterwarnings("ignore", message="Unable to read chart.*", category=UserWarning)
            workbook = self._load_template_workbook(template_path)
        target_sheet_name = analysis_plan.template_sheet_name or self.business_rules.get("template", {}).get("default_sheet_name")
        if target_sheet_name not in workbook.sheetnames:
            raise ValueError(f"模板中不存在目标 sheet：{target_sheet_name}")
        worksheet = workbook[target_sheet_name]

        header_row, data_start_row = self._detect_template_rows(worksheet)
        column_mapping = self._build_template_column_mapping(summary_frame, worksheet, header_row)
        self._normalize_writeable_area(worksheet, min_row=header_row)
        self._clear_template_data_area(worksheet, data_start_row)

        style_row = data_start_row
        for row_offset, (_, row) in enumerate(summary_frame.iterrows(), start=0):
            target_row = data_start_row + row_offset
            for source_index, value in enumerate(row.tolist(), start=1):
                column_index = column_mapping[source_index]
                target_cell = worksheet.cell(row=target_row, column=column_index)
                source_style_cell = worksheet.cell(row=style_row, column=min(column_index, worksheet.max_column))
                self._copy_cell_style(source_style_cell, target_cell)
                target_cell.value = self._coerce_excel_value(value)

        audit_sheet = workbook["审计明细"] if "审计明细" in workbook.sheetnames else workbook.create_sheet("审计明细")
        self._write_simple_frame(audit_sheet, pd.DataFrame(analysis_result.audit_records))

        workbook.save(excel_path)

    def _build_template_column_mapping(self, summary_frame: pd.DataFrame, worksheet, header_row: int) -> dict[int, int]:
        template_headers = {
            column_index: worksheet.cell(row=header_row, column=column_index).value
            for column_index in range(1, worksheet.max_column + 1)
        }
        used_columns: set[int] = set()
        column_mapping: dict[int, int] = {}
        append_column = worksheet.max_column + 1

        for source_index, column_name in enumerate(summary_frame.columns, start=1):
            matched_index = self._find_template_column_index(column_name, template_headers, used_columns)
            if matched_index is None:
                matched_index = append_column
                worksheet.cell(row=header_row, column=matched_index, value=column_name)
                append_column += 1
            used_columns.add(matched_index)
            column_mapping[source_index] = matched_index

        return column_mapping

    def _find_template_column_index(
        self,
        summary_column: str,
        template_headers: dict[int, Any],
        used_columns: set[int],
    ) -> int | None:
        normalized_summary = normalize_text(summary_column)
        alias_groups = {
            "阶段": ["phase", "stage", "阶段"],
            "测试项目": ["item", "testitem", "测试项目", "测试项"],
            "编号": ["no", "no.", "编号"],
            "SN号": ["sn", "sn号", "serial"],
            "测试节点": ["status", "checkpoint", "testnode", "节点", "测试节点"],
            "remark": ["remark", "备注"],
        }

        if summary_column in alias_groups:
            return self._match_template_aliases(alias_groups[summary_column], template_headers, used_columns)

        if summary_column.endswith(" 原始值"):
            metric_name = summary_column[: -len(" 原始值")]
            return self._match_metric_template_column(metric_name, template_headers, used_columns, decay=False)
        if summary_column.endswith(" 衰减"):
            metric_name = summary_column[: -len(" 衰减")]
            return self._match_metric_template_column(metric_name, template_headers, used_columns, decay=True)

        for column_index, template_header in template_headers.items():
            if column_index in used_columns:
                continue
            if normalize_text(template_header) == normalized_summary:
                return column_index
        return None

    def _match_template_aliases(
        self,
        aliases: list[str],
        template_headers: dict[int, Any],
        used_columns: set[int],
    ) -> int | None:
        normalized_aliases = [normalize_text(alias) for alias in aliases]
        for column_index, template_header in template_headers.items():
            if column_index in used_columns:
                continue
            normalized_header = normalize_text(template_header)
            if normalized_header in normalized_aliases:
                return column_index
        return None

    def _match_metric_template_column(
        self,
        metric_name: str,
        template_headers: dict[int, Any],
        used_columns: set[int],
        *,
        decay: bool,
    ) -> int | None:
        normalized_metric = normalize_text(metric_name)
        for column_index, template_header in template_headers.items():
            if column_index in used_columns:
                continue
            normalized_header = normalize_text(template_header)
            if normalized_metric not in normalized_header:
                continue
            header_is_decay = any(token in normalized_header for token in ["decreasing", "decreased", "衰减", "decay"])
            if decay == header_is_decay:
                return column_index
        return None

    def _load_template_workbook(self, template_path: Path):
        try:
            return load_workbook(template_path)
        except TypeError as exc:
            if "expected <class 'str'>" not in str(exc):
                raise
            sanitized_copy = self._build_chartless_template_copy(template_path)
            return load_workbook(sanitized_copy)

    def _build_chartless_template_copy(self, template_path: Path) -> Path:
        with NamedTemporaryFile(prefix="template_sanitized_", suffix=".xlsx", delete=False) as tmp:
            sanitized_path = Path(tmp.name)

        with ZipFile(template_path, "r") as source_zip, ZipFile(sanitized_path, "w", ZIP_DEFLATED) as target_zip:
            for info in source_zip.infolist():
                file_name = info.filename
                if file_name.startswith(("xl/drawings/", "xl/charts/", "xl/media/")):
                    continue

                content = source_zip.read(file_name)
                if file_name.startswith("xl/worksheets/") and file_name.endswith(".xml"):
                    text = content.decode("utf-8", errors="ignore")
                    text = re.sub(r"<drawing[^>]*/>", "", text)
                    text = re.sub(r"<legacyDrawing[^>]*/>", "", text)
                    text = re.sub(r"<picture[^>]*/>", "", text)
                    content = text.encode("utf-8")
                elif file_name.startswith("xl/worksheets/_rels/") and file_name.endswith(".rels"):
                    text = content.decode("utf-8", errors="ignore")
                    text = re.sub(r'<Relationship[^>]+Target="\.\./drawings/[^"]+"[^>]*/>', "", text)
                    text = re.sub(r'<Relationship[^>]+Type="[^"]+/drawing"[^>]*/>', "", text)
                    text = re.sub(r'<Relationship[^>]+Type="[^"]+/chart"[^>]*/>', "", text)
                    content = text.encode("utf-8")

                target_zip.writestr(info, content)
        return sanitized_path

    def _detect_template_rows(self, worksheet) -> tuple[int, int]:
        search_rows = self.business_rules.get("template", {}).get("detect_header_within_top_rows", 20)
        best_row = 1
        best_non_empty = 0
        for row_index in range(1, min(search_rows, worksheet.max_row) + 1):
            non_empty = sum(
                1
                for cell in worksheet[row_index]
                if cell.value not in (None, "")
            )
            if non_empty > best_non_empty:
                best_non_empty = non_empty
                best_row = row_index
        return best_row, best_row + 1

    def _normalize_writeable_area(self, worksheet, min_row: int) -> None:
        ranges_to_unmerge = [
            str(cell_range)
            for cell_range in worksheet.merged_cells.ranges
            if cell_range.max_row >= min_row
        ]
        for cell_range in ranges_to_unmerge:
            worksheet.unmerge_cells(cell_range)

    def _clear_template_data_area(self, worksheet, data_start_row: int) -> None:
        for row in worksheet.iter_rows(min_row=data_start_row, max_row=worksheet.max_row):
            for cell in row:
                cell.value = None

    def _copy_cell_style(self, source_cell, target_cell) -> None:
        if source_cell.has_style:
            target_cell.font = copy(source_cell.font)
            target_cell.border = copy(source_cell.border)
            target_cell.fill = copy(source_cell.fill)
            target_cell.number_format = copy(source_cell.number_format)
            target_cell.protection = copy(source_cell.protection)
            target_cell.alignment = copy(source_cell.alignment)

    def _coerce_excel_value(self, value: Any) -> Any:
        if isinstance(value, (dict, list, tuple, set)):
            return str(jsonify(value))
        if pd.isna(value):
            return None
        return value

    def _write_simple_frame(self, worksheet, frame: pd.DataFrame) -> None:
        worksheet.delete_rows(1, worksheet.max_row)
        if frame.empty:
            worksheet["A1"] = "暂无审计记录"
            return
        for column_index, column_name in enumerate(frame.columns, start=1):
            worksheet.cell(row=1, column=column_index, value=column_name)
        for row_offset, (_, row) in enumerate(frame.iterrows(), start=2):
            for column_index, value in enumerate(row.tolist(), start=1):
                worksheet.cell(row=row_offset, column=column_index, value=self._coerce_excel_value(value))

    def _write_charts(self, job_dir: Path, analysis_result: AnalysisResult) -> list[Path]:
        summary_frame = analysis_result.summary_frame
        assert summary_frame is not None

        chart_paths: list[Path] = []
        x_axis = "测试节点" if "测试节点" in summary_frame.columns else "编号"
        color_axis = "来源文件" if "来源文件" in summary_frame.columns else ("SN号" if "SN号" in summary_frame.columns else None)

        for match in analysis_result.metric_matches:
            raw_column = f"{match.requested_name} 原始值"
            if raw_column not in summary_frame.columns:
                continue

            numeric_values = pd.to_numeric(summary_frame[raw_column], errors="coerce")
            if numeric_values.notna().sum() == 0:
                continue

            figure = px.line(
                summary_frame.assign(**{raw_column: numeric_values}),
                x=x_axis,
                y=raw_column,
                color=color_axis,
                markers=True,
                title=f"{match.requested_name} 趋势图",
            )
            analysis_result.chart_figures[match.requested_name] = figure

            chart_path = job_dir / f"{safe_filename(match.requested_name)}.html"
            figure.write_html(str(chart_path))
            chart_paths.append(chart_path)

        return chart_paths

    def _build_markdown_report(
        self,
        workbook: WorkbookContext,
        analysis_plan: AnalysisPlan,
        analysis_result: AnalysisResult,
        artifacts: AnalysisArtifacts,
    ) -> str:
        lines: list[str] = [
            "# 测试数据分析报告",
            "",
            "## 1. 输入概览",
            "",
            f"- 输入文件：{', '.join(analysis_result.source_files) if analysis_result.source_files else (workbook.file_path.name if workbook.file_path else '未识别')}",
            f"- 数据源 sheet：`{analysis_result.source_sheet}`",
            f"- 请求测试项：{', '.join(target.label for target in analysis_plan.metric_targets) if analysis_plan.metric_targets else '未指定'}",
            f"- 分组维度：{', '.join(analysis_plan.group_dimensions) if analysis_plan.group_dimensions else '未指定'}",
            f"- 统计项：{', '.join(analysis_plan.statistics) if analysis_plan.statistics else '未指定'}",
            f"- 衰减公式：`{analysis_plan.decay_method or '未指定'}`",
            f"- 公式说明：{DECAY_METHODS.get(analysis_plan.decay_method or '', '未指定')}",
            f"- 模板文件：`{workbook.template_file.name if workbook.template_file else analysis_plan.template_hint or '未提供'}`",
            "",
            "## 2. 字段映射",
            "",
        ]

        for key, value in analysis_result.base_field_mapping.items():
            lines.append(f"- `{key}` -> `{value}`")
        if not analysis_result.base_field_mapping:
            lines.append("- 未检测到基础字段映射。")

        lines.extend(["", "## 3. 指标匹配结果", ""])
        for match in analysis_result.metric_matches:
            lines.append(
                f"- `{match.requested_name}` -> `{match.matched_column or '未匹配'}` "
                f"(status={match.status}, score={match.score}, reason={match.reason})"
            )

        lines.extend(["", "## 4. 统计摘要", ""])
        if analysis_result.statistics:
            for metric, stats in analysis_result.statistics.items():
                lines.append(
                    f"- `{metric}`: count={stats.get('count')}, mean={stats.get('mean')}, "
                    f"min={stats.get('min')}, max={stats.get('max')}, std={stats.get('std')}"
                )
        else:
            lines.append("- 当前没有可计算的数值统计。")

        lines.extend(["", "## 5. 智能体结论摘要", ""])
        if analysis_result.narrative_summary:
            for item in analysis_result.narrative_summary:
                lines.append(f"- {item}")
        else:
            lines.append("- 当前没有 LLM 生成的结论摘要。")

        lines.extend(["", "## 6. 多文件/版本说明", ""])
        if analysis_result.comparison_notes:
            for note in analysis_result.comparison_notes:
                lines.append(f"- {note}")
        else:
            lines.append("- 当前任务未触发多文件对比说明。")

        lines.extend(["", "## 7. 风险与告警", ""])
        for message in analysis_result.validation_messages + analysis_result.warnings:
            lines.append(f"- {message}")
        if not analysis_result.validation_messages and not analysis_result.warnings:
            lines.append("- 无。")

        lines.extend(["", "## 8. 模型与技能备注", ""])
        if analysis_result.llm_notes:
            for item in analysis_result.llm_notes:
                lines.append(f"- {item}")
        else:
            lines.append("- 无。")

        lines.extend(
            [
                "",
                "## 9. 产物路径",
                "",
                f"- Excel：`{artifacts.excel_path.name if artifacts.excel_path else '未生成'}`",
                f"- Markdown：`{artifacts.markdown_path.name if artifacts.markdown_path else '未生成'}`",
                f"- PDF：`{artifacts.pdf_path.name if artifacts.pdf_path else '未生成'}`",
                f"- Docx：`{artifacts.docx_path.name if artifacts.docx_path else '未生成'}`",
                f"- 审计日志：`{artifacts.audit_log_path.name if artifacts.audit_log_path else '未生成'}`",
            ]
        )
        if artifacts.output_dir:
            lines.append(f"- 输出目录：`{artifacts.output_dir}`")
        return "\n".join(lines) + "\n"

    def _build_audit_report(
        self,
        workbook: WorkbookContext,
        analysis_plan: AnalysisPlan,
        analysis_result: AnalysisResult,
        artifacts: AnalysisArtifacts,
    ) -> str:
        lines = [
            "# 审计报告",
            "",
            "## 1. 任务信息",
            "",
            f"- 数据源 sheet：`{analysis_result.source_sheet}`",
            f"- 输入文件数：{len(analysis_result.source_files) or 1}",
            f"- 请求指标数：{len(analysis_plan.metric_targets)}",
            f"- 输出粒度：{analysis_plan.output_granularity}",
            "",
            "## 2. 校验结果",
            "",
        ]
        for message in analysis_result.validation_messages:
            lines.append(f"- {message}")
        if not analysis_result.validation_messages:
            lines.append("- 无。")

        lines.extend(["", "## 3. 审计明细", ""])
        for record in analysis_result.audit_records:
            lines.append(f"- {record}")
        if not analysis_result.audit_records:
            lines.append("- 无审计记录。")

        lines.extend(["", "## 4. 可复算复核", ""])
        for review in analysis_result.recompute_reviews:
            lines.append(f"- {review}")
        if not analysis_result.recompute_reviews:
            lines.append("- 无。")

        lines.extend(["", "## 5. 智能体总结", ""])
        for item in analysis_result.narrative_summary:
            lines.append(f"- {item}")
        if not analysis_result.narrative_summary:
            lines.append("- 无。")

        lines.extend(["", "## 6. Agent Trace", ""])
        for event in analysis_result.trace_events:
            lines.append(
                f"- [{event.agent}] {event.step} | status={event.status} | duration_ms={event.duration_ms} | {event.summary}"
            )
        if not analysis_result.trace_events:
            lines.append("- 当前没有 trace 记录。")
        return "\n".join(lines) + "\n"

    def _write_docx(self, docx_path: Path, markdown_text: str) -> bool:
        if Document is None:
            return False
        document = Document()
        for line in markdown_text.splitlines():
            if line.startswith("# "):
                document.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                document.add_heading(line[3:], level=2)
            elif line.startswith("- "):
                document.add_paragraph(line[2:], style="List Bullet")
            elif line.strip():
                document.add_paragraph(line)
        document.save(docx_path)
        return True

    def _write_pdf(self, pdf_path: Path, markdown_text: str) -> bool:
        if canvas is None or pdfmetrics is None or UnicodeCIDFont is None or A4 is None:
            return False
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        page = canvas.Canvas(str(pdf_path), pagesize=A4)
        page.setFont("STSong-Light", 11)
        width, height = A4
        y = height - 40
        for raw_line in markdown_text.splitlines():
            line = raw_line.strip() or " "
            if y < 40:
                page.showPage()
                page.setFont("STSong-Light", 11)
                y = height - 40
            page.drawString(40, y, line[:95])
            y -= 16
        page.save()
        return True
